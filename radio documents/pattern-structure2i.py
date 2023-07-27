from docx import Document
from docx.shared import RGBColor, Pt
from dateutil.parser import parse
import re
import os

def is_date(string, fuzzy=False):
    """
    Return whether the string can be interpreted as a date.
    """
    try: 
        parse(string, fuzzy=fuzzy)
        return True
    except ValueError:
        return False

def format_received_and_dates(filename):
    doc = Document(filename)
    new_doc = Document()

    for paragraph in doc.paragraphs:
        parts = re.split(r'(received)', paragraph.text, flags=re.IGNORECASE)
        date_found = False
        for part in parts:
            if part:
                new_paragraph = new_doc.add_paragraph()
                words = part.strip().split()
                for word in words:
                    new_run = new_paragraph.add_run(f"{word} ")
                    new_run.font.name = 'Arial'
                    new_run.font.size = Pt(11)

                    if re.match(r'received', word, re.IGNORECASE):
                        new_run.font.color.rgb = RGBColor(255, 0, 0)  # red color
                        new_run.bold = True
                        date_found = True

                    elif date_found and is_date(word):
                        new_run.font.color.rgb = RGBColor(0, 0, 255)  # blue color
                        new_run.bold = True
                        date_found = False  # reset for next "received"

    # Save the document
    new_doc.save('formatted_and_date_highlighted_' + os.path.basename(filename))

format_received_and_dates('place path to document here')
